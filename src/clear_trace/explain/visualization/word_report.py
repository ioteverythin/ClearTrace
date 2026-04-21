"""Generate a single professional Word document (.docx) from Prism analysis.

Produces a unified systems-engineering report that covers:
  - Prompt analysis (if a PromptReport is provided)
  - Tool pipeline analysis (if a ToolReport is provided)
  - Or both combined into one document

Includes:
  - Title page with scores
  - Executive summary
  - Original prompt (before) vs. Improved prompt (after) with highlighted changes
  - All SE matrices as Word tables
  - Matplotlib charts for BOTH prompt and tool matrices
  - Detailed findings explaining WHY each change was made
  - Verification summary

Usage:
    >>> from clear_trace import PromptAdvisor, ToolAdvisor, WordReport
    >>> # Prompt only
    >>> WordReport(prompt_report=prompt_report).save("report.docx")
    >>> # Tool only
    >>> WordReport(tool_report=tool_report).save("report.docx")
    >>> # Combined
    >>> WordReport(prompt_report=pr, tool_report=tr).save("report.docx")
"""

from __future__ import annotations

import io
import os
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from clear_trace.explain.advisor.suggestions import PromptReport, Suggestion, ImpactLevel
from clear_trace.explain.advisor.matrix_report import MatrixReport, Matrix
from clear_trace.explain.advisor.tool_types import ToolReport, ToolDefinition, Severity
from clear_trace.explain.advisor.tool_matrix_report import ToolMatrixReport


# ── Colors ───────────────────────────────────────────────────────

_BLUE = RGBColor(0x1A, 0x73, 0xE8)
_DARK = RGBColor(0x20, 0x24, 0x2A)
_GREEN = RGBColor(0x34, 0xA8, 0x53)
_RED = RGBColor(0xEA, 0x43, 0x35)
_GRAY = RGBColor(0x5F, 0x63, 0x68)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_HDR_BG = "1A73E8"
_BG_GREEN = "D4EDDA"
_BG_RED = "F8D7DA"
_BG_YELLOW = "FFF3CD"
_BG_BLUE = "D1ECF1"


# ── Helpers ──────────────────────────────────────────────────────

def _cell_bg(cell, hex_color: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _status_bg(val: str) -> str:
    s = str(val).upper()
    if s in ("PASS", "OK", "COVERED", "FIXED", "IMPROVED", "YES"):
        return _BG_GREEN
    if s in ("FAIL", "GAP", "REGRESSED", "CRITICAL", "NO"):
        return _BG_RED
    if s in ("WARN", "WARNING", "MARGINAL", "PARTIAL", "SAME", "INFO"):
        return _BG_YELLOW
    return "FFFFFF"


def _pct(v: float) -> str:
    return f"{v * 100:.0f}%" if v <= 1.0 else f"{v:.0f}%"


def _table(doc: Document, headers: List[str], rows: List[Dict[str, str]],
           color_col: Optional[str] = None):
    """Insert a formatted Word table."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        _cell_bg(c, _HDR_BG)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.color.rgb = _WHITE
                r.font.bold = True
                r.font.size = Pt(9)

    for ri, rd in enumerate(rows):
        for ci, h in enumerate(headers):
            c = t.rows[ri + 1].cells[ci]
            v = str(rd.get(h, ""))
            c.text = v
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8)
            if color_col and h == color_col:
                _cell_bg(c, _status_bg(v))
    return t


def _embed_fig(doc: Document, fig, width: float = 7.5):
    """Render matplotlib figure into the doc as an image."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    buf.seek(0)
    doc.add_picture(buf, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    import matplotlib.pyplot as plt
    plt.close(fig)


def _find_color_col(headers: List[str]) -> Optional[str]:
    """Pick the best column to color-code."""
    for c in ("Status", "Assessment", "Severity", "Risk Level",
              "Output Status", "Match", "JSON Valid"):
        if c in headers:
            return c
    return None


def _write_diff_cell(cell, before_text: str, after_text: str):
    """Write a unified diff into a Word cell with color-coded lines.

    - Unchanged lines: normal black text
    - Removed lines:   red text with strikethrough
    - Added lines:     green bold text with highlight
    """
    import difflib

    before_lines = before_text.splitlines(keepends=True)
    after_lines = after_text.splitlines(keepends=True)

    diff = list(difflib.unified_diff(
        before_lines, after_lines, lineterm="", n=1,
    ))

    # Clear existing paragraph
    for p in cell.paragraphs:
        p.clear()

    para = cell.paragraphs[0]
    if not diff:
        # No difference
        run = para.add_run("(No changes detected)")
        run.font.size = Pt(9)
        run.italic = True
        return

    # Skip the --- / +++ / @@ headers
    change_count = 0
    for line in diff:
        if line.startswith("---") or line.startswith("+++"):
            continue
        if line.startswith("@@"):
            if change_count > 0:
                run = para.add_run("\n\u2500" * 40 + "\n")
                run.font.size = Pt(7)
                run.font.color.rgb = _GRAY
            change_count += 1
            continue

        text = line.rstrip("\n\r")
        if line.startswith("-"):
            # Removed line — red strikethrough
            run = para.add_run(text + "\n")
            run.font.size = Pt(8)
            run.font.color.rgb = _RED
            run.font.strike = True
        elif line.startswith("+"):
            # Added line — green bold
            run = para.add_run(text + "\n")
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x15, 0x6B, 0x2A)
            run.bold = True
        else:
            # Context line — gray
            run = para.add_run(text + "\n")
            run.font.size = Pt(8)
            run.font.color.rgb = _GRAY


# ── Tool Chart Generator ────────────────────────────────────────

def _make_tool_charts(tmx: ToolMatrixReport) -> List:
    """Generate matplotlib figures for tool matrices."""
    try:
        import matplotlib.pyplot as plt
        import numpy as np
    except ImportError:
        return []

    figs: list = []
    bg = "#ffffff"
    text_c = "#222222"
    title_c = "#0066cc"

    # ── T1: Schema Quality bar chart ─────────────────────────
    t1 = tmx.t1
    if t1.rows:
        fig, ax = plt.subplots(figsize=(10, max(3, len(t1.rows) * 0.4)))
        fig.patch.set_facecolor(bg)
        ax.set_facecolor(bg)

        labels = [
            f"{r.cells.get('Tool', '')}.{r.cells.get('Parameter', '')}"
            for r in t1.rows
        ]
        labels = [la[:30] for la in labels]
        statuses = [r.cells.get("Status", "") for r in t1.rows]
        colors = [
            "#34A853" if s == "PASS" else "#EA4335" if s == "FAIL" else "#FBBC04"
            for s in statuses
        ]

        y = range(len(labels))
        issues = []
        for r in t1.rows:
            try:
                issues.append(int(r.cells.get("Issues", "0")))
            except (ValueError, TypeError):
                issues.append(0)

        ax.barh(y, issues, color=colors, edgecolor="#dddddd")
        ax.set_yticks(list(y))
        ax.set_yticklabels(labels, fontsize=7, color=text_c)
        ax.set_xlabel("Issues", color=text_c)
        ax.set_title(
            "T1: Tool Schema Quality \u2014 Issues per Parameter",
            fontsize=12, fontweight="bold", color=title_c,
        )
        ax.invert_yaxis()
        fig.tight_layout()
        figs.append(("Schema Quality", fig))

    # ── T2: Prompt-Tool Alignment heatmap ────────────────────
    t2 = tmx.t2
    if t2.rows:
        check_cols = [
            "Mentioned in Prompt", "Usage Guidance", "JSON Format",
            "Fallback", "Disambiguation",
        ]
        present_cols = [c for c in check_cols if c in t2.headers]
        if present_cols:
            tools = [r.cells.get("Tool", "") for r in t2.rows]
            data = []
            for r in t2.rows:
                row_vals = []
                for c in present_cols:
                    v = r.cells.get(c, "").upper()
                    if v in ("YES",):
                        row_vals.append(1.0)
                    elif v in ("PARTIAL",):
                        row_vals.append(0.5)
                    else:
                        row_vals.append(0.0)
                data.append(row_vals)

            fig, ax = plt.subplots(figsize=(8, max(2.5, len(tools) * 0.6)))
            fig.patch.set_facecolor(bg)
            ax.set_facecolor(bg)

            arr = np.array(data)
            from matplotlib.colors import LinearSegmentedColormap
            cmap = LinearSegmentedColormap.from_list(
                "rg", ["#EA4335", "#FBBC04", "#34A853"],
            )
            im = ax.imshow(arr, cmap=cmap, aspect="auto", vmin=0, vmax=1)

            ax.set_xticks(range(len(present_cols)))
            ax.set_xticklabels(
                present_cols, fontsize=8, rotation=30, ha="right", color=text_c,
            )
            ax.set_yticks(range(len(tools)))
            ax.set_yticklabels(tools, fontsize=9, color=text_c)

            for ti in range(len(tools)):
                for tj in range(len(present_cols)):
                    val = arr[ti, tj]
                    label = (
                        "YES" if val == 1 else "PARTIAL" if val == 0.5 else "NO"
                    )
                    ax.text(
                        tj, ti, label, ha="center", va="center", fontsize=8,
                        color="white" if val < 0.5 else "black",
                    )

            ax.set_title(
                "T2: Prompt\u2013Tool Alignment",
                fontsize=12, fontweight="bold", color=title_c,
            )
            fig.tight_layout()
            figs.append(("Prompt-Tool Alignment", fig))

    # ── T3: Test accuracy chart ──────────────────────────────
    t3 = tmx.t3
    if t3.rows:
        fig, ax = plt.subplots(figsize=(10, max(3, len(t3.rows) * 0.5)))
        fig.patch.set_facecolor(bg)
        ax.set_facecolor(bg)

        scenarios = [
            r.cells.get("Test Scenario", "")[:30] for r in t3.rows
        ]

        y = range(len(scenarios))
        tool_ok = [
            1 if r.cells.get("Tool Correct") == "YES" else 0 for r in t3.rows
        ]
        param_ok = [
            1 if r.cells.get("Params Correct") == "YES" else 0
            for r in t3.rows
        ]
        json_ok = [
            1 if r.cells.get("JSON Valid") == "YES" else 0 for r in t3.rows
        ]

        ax.barh(y, tool_ok, 0.25, label="Tool Correct", color="#34A853")
        ax.barh(
            [i + 0.25 for i in y], param_ok, 0.25,
            label="Params Correct", color="#4285F4",
        )
        ax.barh(
            [i + 0.5 for i in y], json_ok, 0.25,
            label="JSON Valid", color="#FBBC04",
        )

        ax.set_yticks([i + 0.25 for i in y])
        ax.set_yticklabels(scenarios, fontsize=8, color=text_c)
        ax.set_xlabel("Pass (1) / Fail (0)", color=text_c)
        ax.set_title(
            "T3: Tool Call Accuracy by Test Case",
            fontsize=12, fontweight="bold", color=title_c,
        )
        ax.legend(fontsize=8, loc="lower right")
        ax.invert_yaxis()
        fig.tight_layout()
        figs.append(("Test Accuracy", fig))

    # ── T5: Severity bar ─────────────────────────────────────
    t5 = tmx.t5
    if t5.rows:
        sev_counts: Dict[str, int] = {"CRITICAL": 0, "WARNING": 0, "INFO": 0}
        for r in t5.rows:
            s = r.cells.get("Severity", "").upper()
            if s in sev_counts:
                sev_counts[s] += 1

        fig, ax = plt.subplots(figsize=(6, 3))
        fig.patch.set_facecolor(bg)
        ax.set_facecolor(bg)

        labels = list(sev_counts.keys())
        vals = list(sev_counts.values())
        colors = ["#EA4335", "#FBBC04", "#4285F4"]
        ax.bar(labels, vals, color=colors, edgecolor="#dddddd")
        ax.set_ylabel("Count", color=text_c)
        ax.set_title(
            "T5: Issues by Severity",
            fontsize=12, fontweight="bold", color=title_c,
        )

        for idx, v in enumerate(vals):
            if v > 0:
                ax.text(
                    idx, v + 0.1, str(v), ha="center", fontsize=10,
                    fontweight="bold", color=text_c,
                )

        fig.tight_layout()
        figs.append(("Issue Severity", fig))

    return figs


# =====================================================================
#  WordReport  --  Unified single-document report
# =====================================================================

class WordReport:
    """Generate a single Word document from Prism analysis results.

    Accepts a PromptReport, a ToolReport, or both.  When both are
    provided the report merges them into one document with prompt
    sections first, then tool sections, followed by shared
    verification and appendix sections.

    Args:
        prompt_report: PromptReport from PromptAdvisor.analyze() (optional).
        tool_report: ToolReport from ToolAdvisor.analyze() (optional).
        include_charts: Embed matplotlib charts in the document.
        title: Custom report title.
    """

    def __init__(
        self,
        prompt_report: Optional[PromptReport] = None,
        tool_report: Optional[ToolReport] = None,
        include_charts: bool = True,
        title: str = "Prism Pipeline Analysis Report",
    ):
        if not prompt_report and not tool_report:
            raise ValueError(
                "Provide at least one of prompt_report or tool_report"
            )

        self.pr = prompt_report
        self.tr = tool_report
        self.charts = include_charts
        self.title = title

        self._prompt_mx = MatrixReport(prompt_report) if prompt_report else None
        self._tool_mx = ToolMatrixReport(tool_report) if tool_report else None

    # ------------------------------------------------------------------
    #  Public API
    # ------------------------------------------------------------------

    def save(self, path: str) -> str:
        """Build and save the Word document. Returns absolute path."""
        os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
        doc = Document()

        # Landscape orientation
        sec = doc.sections[0]
        sec.page_width = Inches(11)
        sec.page_height = Inches(8.5)
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.left_margin = Cm(2)
        sec.right_margin = Cm(2)
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)

        _sec = [0]  # mutable counter

        def _h(title: str) -> str:
            _sec[0] += 1
            return f"{_sec[0]}. {title}"

        # ── Build Document ──────────────────────────────────
        self._title_page(doc)

        doc.add_heading(_h("Executive Summary"), level=1)
        self._executive_summary(doc)
        doc.add_page_break()

        # -- Prompt sections --
        if self.pr:
            doc.add_heading(_h("Prompt \u2014 Before & After"), level=1)
            self._prompt_before_after(doc)
            doc.add_page_break()

            doc.add_heading(_h("Prompt \u2014 SE Matrices"), level=1)
            self._prompt_matrices(doc)
            doc.add_page_break()

            if self.charts:
                doc.add_heading(_h("Prompt \u2014 Charts"), level=1)
                self._prompt_charts(doc)
                doc.add_page_break()

            doc.add_heading(
                _h("Prompt \u2014 Why Each Change Was Made"), level=1,
            )
            self._prompt_findings(doc)
            doc.add_page_break()

        # -- Tool sections --
        if self.tr:
            doc.add_heading(_h("Tool Pipeline \u2014 Overview"), level=1)
            self._tool_overview(doc)
            doc.add_page_break()

            doc.add_heading(
                _h("Tool Pipeline \u2014 SE Matrices"), level=1,
            )
            self._tool_matrices(doc)
            doc.add_page_break()

            if self.charts:
                doc.add_heading(
                    _h("Tool Pipeline \u2014 Charts"), level=1,
                )
                self._tool_charts(doc)
                doc.add_page_break()

            doc.add_heading(
                _h("Tool Pipeline \u2014 Findings & Rationale"), level=1,
            )
            self._tool_findings(doc)
            doc.add_page_break()

            doc.add_heading(
                _h("Tool Pipeline \u2014 Improvements"), level=1,
            )
            self._tool_improvements(doc)
            doc.add_page_break()

        # -- Shared closing sections --
        doc.add_heading(_h("Verification Summary"), level=1)
        self._verification(doc)
        doc.add_page_break()

        doc.add_heading(
            "Appendix: Methodology & Matrix Definitions", level=1,
        )
        self._appendix(doc)

        doc.save(path)
        return os.path.abspath(path)

    # ==================================================================
    #  Title page
    # ==================================================================

    def _title_page(self, doc: Document):
        for _ in range(3):
            doc.add_paragraph("")

        h = doc.add_heading("PRISM", level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in h.runs:
            r.font.color.rgb = _BLUE
            r.font.size = Pt(48)

        h2 = doc.add_heading(self.title, level=1)
        h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in h2.runs:
            r.font.color.rgb = _DARK
            r.font.size = Pt(22)

        doc.add_paragraph("")

        # Score lines
        lines: list[str] = []
        if self.pr:
            lines.append(
                f"Prompt Score:  {_pct(self.pr.score_before)}"
                f"  \u2192  {_pct(self.pr.score_after)}"
            )
        if self.tr:
            lines.append(
                f"Tool Score:  {_pct(self.tr.score_before)}"
                f"  \u2192  {_pct(self.tr.score_after)}"
            )

        for line in lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(18)
            run.bold = True
            run.font.color.rgb = _GREEN

        doc.add_paragraph("")
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta.add_run(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        ).font.size = Pt(10)
        if self.pr:
            meta.add_run(
                f"   |   Suggestions: {self.pr.num_suggestions}"
            ).font.size = Pt(10)
        if self.tr:
            meta.add_run(
                f"   |   Tools: {len(self.tr.tools)}"
            ).font.size = Pt(10)
            meta.add_run(
                f"   |   Issues: {len(self.tr.all_issues)}"
            ).font.size = Pt(10)

        doc.add_page_break()

    # ==================================================================
    #  Executive summary
    # ==================================================================

    def _executive_summary(self, doc: Document):
        doc.add_paragraph(
            "This report presents a Prism systems-engineering analysis. "
            "The pipeline was decomposed, tested against the user\u2019s LLM, "
            "diagnosed, and improved. Each finding includes the root cause "
            "(WHY) and the corrective action (WHAT)."
        )

        headers = ["Metric", "Value"]
        rows: list[dict] = []

        if self.pr:
            rows.extend([
                {"Metric": "Prompt Score Before",
                 "Value": _pct(self.pr.score_before)},
                {"Metric": "Prompt Score After",
                 "Value": _pct(self.pr.score_after)},
                {"Metric": "Prompt Improvement",
                 "Value": f"+{_pct(self.pr.score_after - self.pr.score_before)}"},
                {"Metric": "Suggestions",
                 "Value": str(self.pr.num_suggestions)},
            ])
        if self.tr:
            rows.extend([
                {"Metric": "Tool Score Before",
                 "Value": _pct(self.tr.score_before)},
                {"Metric": "Tool Score After",
                 "Value": _pct(self.tr.score_after)},
                {"Metric": "Test Pass Rate",
                 "Value": f"{self.tr.test_pass_rate:.0%}"},
                {"Metric": "Critical Issues",
                 "Value": str(self.tr.critical_count)},
                {"Metric": "Total Tool Issues",
                 "Value": str(len(self.tr.all_issues))},
            ])

        _table(doc, headers, rows)

        if self.pr and self.pr.diagnosis:
            doc.add_paragraph("")
            doc.add_heading("Diagnosis", level=3)
            doc.add_paragraph(self.pr.diagnosis)

    # ==================================================================
    #  Prompt -- Before / After with highlighted changes
    # ==================================================================

    def _prompt_before_after(self, doc: Document):
        doc.add_paragraph(
            "Below is the original prompt alongside the improved prompt "
            "generated by Prism. Changes address the findings detailed "
            "later in this report."
        )

        # Side-by-side table: BEFORE | AFTER
        t = doc.add_table(rows=2, cols=2)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Headers
        _cell_bg(t.rows[0].cells[0], _BG_RED)
        _cell_bg(t.rows[0].cells[1], _BG_GREEN)
        for p in t.rows[0].cells[0].paragraphs:
            r = p.add_run("BEFORE \u2014 Original Prompt")
            r.bold = True
            r.font.size = Pt(11)
        for p in t.rows[0].cells[1].paragraphs:
            r = p.add_run("AFTER \u2014 Improved Prompt")
            r.bold = True
            r.font.size = Pt(11)

        # Content
        before_cell = t.rows[1].cells[0]
        after_cell = t.rows[1].cells[1]

        for p in before_cell.paragraphs:
            r = p.add_run(self.pr.original_prompt)
            r.font.size = Pt(9)

        if self.pr.improved_prompt:
            for p in after_cell.paragraphs:
                r = p.add_run(self.pr.improved_prompt)
                r.font.size = Pt(9)
        else:
            for p in after_cell.paragraphs:
                r = p.add_run("(No improved prompt generated)")
                r.font.size = Pt(9)
                r.italic = True

        for row in t.rows:
            for cell in row.cells:
                cell.width = Inches(4.3)

        doc.add_paragraph("")

        # ── Highlighted diff section ──────────────────────────
        if self.pr.improved_prompt and self.pr.improved_prompt != self.pr.original_prompt:
            doc.add_heading("Changes Highlighted", level=2)
            doc.add_paragraph(
                "Line-by-line diff: "
                "\u0336r\u0336e\u0336d\u0020\u0336s\u0336t\u0336r\u0336i\u0336k\u0336e\u0336t\u0336h\u0336r\u0336o\u0336u\u0336g\u0336h = removed, "
                "green bold = added, gray = unchanged context."
            )

            diff_tbl = doc.add_table(rows=2, cols=1)
            diff_tbl.style = "Table Grid"
            diff_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            _cell_bg(diff_tbl.rows[0].cells[0], _BG_BLUE)
            for p in diff_tbl.rows[0].cells[0].paragraphs:
                r = p.add_run("DIFF \u2014 What Changed")
                r.bold = True
                r.font.size = Pt(11)

            _write_diff_cell(
                diff_tbl.rows[1].cells[0],
                self.pr.original_prompt,
                self.pr.improved_prompt,
            )

        doc.add_paragraph("")

        # Change summary table
        if self.pr.suggestions:
            doc.add_heading("What Changed & Why", level=2)
            doc.add_paragraph(
                f"Prism identified {self.pr.num_suggestions} changes. "
                "Each is explained below with its rationale."
            )

            hdrs = ["#", "Type", "Impact", "What Changed", "Why"]
            rows: list[dict] = []
            for i, s in enumerate(self.pr.suggestions, 1):
                rows.append({
                    "#": str(i),
                    "Type": s.type.value.upper(),
                    "Impact": s.impact.value.upper(),
                    "What Changed": s.fix[:80],
                    "Why": s.problem[:80],
                })
            _table(doc, hdrs, rows, color_col="Impact")

        # Score impact
        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.add_run("Score Impact: ").bold = True
        delta = self.pr.score_after - self.pr.score_before
        sign = "+" if delta >= 0 else ""
        run = p.add_run(
            f"{_pct(self.pr.score_before)} \u2192 "
            f"{_pct(self.pr.score_after)} ({sign}{_pct(delta)})"
        )
        run.font.color.rgb = _GREEN if delta > 0 else _RED if delta < 0 else _GRAY
        run.bold = True

    # ==================================================================
    #  Prompt -- SE Matrices
    # ==================================================================

    def _prompt_matrices(self, doc: Document):
        doc.add_paragraph(
            "Five SE-standard matrices trace each prompt component through "
            "analysis, sensitivity testing, gap analysis, corrective "
            "actions, and verification."
        )

        for mx in self._prompt_mx.matrices:
            doc.add_heading(f"{mx.id}: {mx.title}", level=2)
            p = doc.add_paragraph(mx.description)
            if p.runs:
                p.runs[0].italic = True
            doc.add_paragraph("")
            rows = [r.cells for r in mx.rows]
            _table(doc, mx.headers, rows,
                   color_col=_find_color_col(mx.headers))
            doc.add_paragraph("")

    # ==================================================================
    #  Prompt -- Charts
    # ==================================================================

    def _prompt_charts(self, doc: Document):
        try:
            from clear_trace.explain.visualization.matrix_plots import MatrixPlotter
        except ImportError:
            doc.add_paragraph(
                "matplotlib not available \u2014 charts skipped."
            )
            return

        doc.add_paragraph(
            "Visual representations of the prompt SE matrices. "
            "Green = healthy, Yellow = marginal, Red = requires action."
        )

        plotter = MatrixPlotter(self._prompt_mx, style="light")
        chart_methods = [
            (plotter.plot_pcam, "Prompt Component Assessment"),
            (plotter.plot_srm, "Sensitivity & Robustness"),
            (plotter.plot_rgam, "Requirements Gap Analysis"),
            (plotter.plot_car, "Corrective Action Summary"),
            (plotter.plot_bavm, "Before / After Verification"),
        ]
        for method, caption in chart_methods:
            try:
                fig = method()
                if fig and len(fig.axes) > 0:
                    doc.add_heading(caption, level=3)
                    _embed_fig(doc, fig, width=8.5)
                    doc.add_paragraph("")
            except Exception:
                pass

    # ==================================================================
    #  Prompt -- Findings detail
    # ==================================================================

    def _prompt_findings(self, doc: Document):
        doc.add_paragraph(
            "Detailed rationale for each change. Every finding explains "
            "the root cause (WHY the current prompt fails) and the "
            "corrective action (WHAT was changed)."
        )

        if not self.pr.suggestions:
            doc.add_paragraph("No suggestions were generated.")
            return

        for i, s in enumerate(self.pr.suggestions, 1):
            doc.add_heading(
                f"Finding #{i}:  {s.type.value.upper()} "
                f"\u2014 {s.target[:50]}",
                level=3,
            )

            p = doc.add_paragraph()
            p.add_run("WHY \u2014 Problem: ").bold = True
            p.add_run(s.problem)

            p = doc.add_paragraph()
            p.add_run("WHAT \u2014 Corrective Action: ").bold = True
            p.add_run(s.fix)

            if s.improved_text:
                p = doc.add_paragraph()
                p.add_run("New Text: ").bold = True
                run = p.add_run(f'"{s.improved_text}"')
                run.italic = True

            if s.evidence:
                p = doc.add_paragraph()
                p.add_run("Evidence: ").bold = True
                p.add_run(s.evidence)

            p = doc.add_paragraph()
            run = p.add_run(
                f"Confidence: {s.confidence:.0%}  |  "
                f"Impact: {s.impact.value.upper()}"
            )
            run.font.color.rgb = _GRAY
            run.font.size = Pt(8)

            doc.add_paragraph("")

    # ==================================================================
    #  Tool -- Overview
    # ==================================================================

    def _tool_overview(self, doc: Document):
        doc.add_paragraph(
            "The tool pipeline was analyzed for schema quality, "
            "prompt-tool alignment, live tool call accuracy, and "
            "JSON compliance."
        )

        # System prompt
        doc.add_heading("System Prompt", level=2)
        p = doc.add_paragraph()
        run = p.add_run(self.tr.system_prompt)
        run.font.size = Pt(9)
        run.italic = True

        # Tool definitions
        doc.add_heading("Tool Definitions", level=2)
        for tool in self.tr.tools:
            doc.add_heading(f"{tool.name}", level=3)
            p = doc.add_paragraph()
            p.add_run("Description: ").bold = True
            p.add_run(tool.description or "(missing)")

            props = tool.parameters.get("properties", {})
            required = tool.parameters.get("required", tool.required)
            if props:
                hdrs = ["Parameter", "Type", "Required", "Description"]
                tbl_rows: list[dict] = []
                for pn, pd in props.items():
                    if not isinstance(pd, dict):
                        continue
                    tbl_rows.append({
                        "Parameter": pn,
                        "Type": pd.get("type", "\u2014"),
                        "Required": "YES" if pn in required else "no",
                        "Description": pd.get("description", "(none)")[:60],
                    })
                if tbl_rows:
                    _table(doc, hdrs, tbl_rows)
                    doc.add_paragraph("")

    # ==================================================================
    #  Tool -- SE Matrices
    # ==================================================================

    def _tool_matrices(self, doc: Document):
        doc.add_paragraph(
            "Five SE matrices covering schema quality, prompt alignment, "
            "test accuracy, JSON compliance, and corrective actions."
        )

        for mx in self._tool_mx.matrices:
            doc.add_heading(f"{mx.id}: {mx.title}", level=2)
            p = doc.add_paragraph(mx.description)
            if p.runs:
                p.runs[0].italic = True
            doc.add_paragraph("")
            rows = [r.cells for r in mx.rows]
            _table(doc, mx.headers, rows,
                   color_col=_find_color_col(mx.headers))
            doc.add_paragraph("")

    # ==================================================================
    #  Tool -- Charts  (NEW -- was missing before!)
    # ==================================================================

    def _tool_charts(self, doc: Document):
        doc.add_paragraph(
            "Visual analysis of tool pipeline health. "
            "Each chart maps to an SE matrix from the previous section."
        )

        chart_figs = _make_tool_charts(self._tool_mx)
        if not chart_figs:
            doc.add_paragraph(
                "matplotlib not available \u2014 charts skipped."
            )
            return

        for caption, fig in chart_figs:
            doc.add_heading(caption, level=3)
            _embed_fig(doc, fig, width=8.5)
            doc.add_paragraph("")

    # ==================================================================
    #  Tool -- Findings
    # ==================================================================

    def _tool_findings(self, doc: Document):
        doc.add_paragraph(
            "Every issue found, with severity, root cause (WHY), "
            "corrective action (WHAT), and evidence."
        )

        issues = self.tr.all_issues
        if not issues:
            doc.add_paragraph("No issues found.")
            return

        # Severity summary table
        critical = [i for i in issues if i.severity == Severity.CRITICAL]
        warnings = [i for i in issues if i.severity == Severity.WARNING]
        info = [i for i in issues if i.severity == Severity.INFO]

        _table(doc, ["Severity", "Count"], [
            {"Severity": "CRITICAL", "Count": str(len(critical))},
            {"Severity": "WARNING", "Count": str(len(warnings))},
            {"Severity": "INFO", "Count": str(len(info))},
        ], color_col="Severity")

        doc.add_paragraph("")

        # Issue details
        for idx, issue in enumerate(issues, 1):
            sev = (
                issue.severity.value.upper()
                if hasattr(issue.severity, "value")
                else str(issue.severity).upper()
            )
            doc.add_heading(
                f"CA-{idx:03d}  [{sev}]  {issue.component[:50]}", level=3,
            )

            p = doc.add_paragraph()
            p.add_run("WHY \u2014 Problem: ").bold = True
            p.add_run(issue.problem)

            p = doc.add_paragraph()
            p.add_run("WHAT \u2014 Fix: ").bold = True
            p.add_run(issue.fix)

            if issue.evidence:
                p = doc.add_paragraph()
                p.add_run("Evidence: ").bold = True
                p.add_run(issue.evidence[:200])

    # ==================================================================
    #  Tool -- Improvements  (prompt + schemas before/after)
    # ==================================================================

    def _tool_improvements(self, doc: Document):
        # Improved system prompt: before / after
        doc.add_heading("Improved System Prompt", level=2)

        t = doc.add_table(rows=2, cols=2)
        t.style = "Table Grid"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER

        _cell_bg(t.rows[0].cells[0], _BG_RED)
        _cell_bg(t.rows[0].cells[1], _BG_GREEN)
        for p in t.rows[0].cells[0].paragraphs:
            r = p.add_run("BEFORE \u2014 Original System Prompt")
            r.bold = True
            r.font.size = Pt(11)
        for p in t.rows[0].cells[1].paragraphs:
            r = p.add_run("AFTER \u2014 Improved System Prompt")
            r.bold = True
            r.font.size = Pt(11)

        for p in t.rows[1].cells[0].paragraphs:
            r = p.add_run(self.tr.system_prompt)
            r.font.size = Pt(9)

        after_text = self.tr.improved_prompt or "(No improved prompt generated)"
        for p in t.rows[1].cells[1].paragraphs:
            r = p.add_run(after_text)
            r.font.size = Pt(9)

        for row in t.rows:
            for cell in row.cells:
                cell.width = Inches(4.3)

        doc.add_paragraph("")

        # ── Highlighted diff for tool system prompt ───────────
        if self.tr.improved_prompt and self.tr.improved_prompt != self.tr.system_prompt:
            doc.add_heading("System Prompt Changes Highlighted", level=2)
            doc.add_paragraph(
                "Line-by-line diff: red strikethrough = removed, "
                "green bold = added, gray = unchanged context."
            )

            diff_tbl = doc.add_table(rows=2, cols=1)
            diff_tbl.style = "Table Grid"
            diff_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            _cell_bg(diff_tbl.rows[0].cells[0], _BG_BLUE)
            for p in diff_tbl.rows[0].cells[0].paragraphs:
                r = p.add_run("DIFF \u2014 System Prompt Changes")
                r.bold = True
                r.font.size = Pt(11)

            _write_diff_cell(
                diff_tbl.rows[1].cells[0],
                self.tr.system_prompt,
                self.tr.improved_prompt,
            )
            doc.add_paragraph("")

        # Improved tool schemas
        doc.add_heading("Improved Tool Definitions", level=2)
        if self.tr.improved_tools:
            for tool in self.tr.improved_tools:
                doc.add_heading(f"{tool.name}", level=3)
                p = doc.add_paragraph()
                p.add_run("Description: ").bold = True
                p.add_run(tool.description or "(none)")

                props = tool.parameters.get("properties", {})
                required = tool.parameters.get("required", tool.required)
                if props:
                    hdrs = ["Parameter", "Type", "Required", "Description"]
                    tbl_rows: list[dict] = []
                    for pn, pd in props.items():
                        if not isinstance(pd, dict):
                            continue
                        tbl_rows.append({
                            "Parameter": pn,
                            "Type": pd.get("type", "\u2014"),
                            "Required": "YES" if pn in required else "no",
                            "Description": (
                                pd.get("description", "(none)")[:60]
                            ),
                        })
                    if tbl_rows:
                        _table(doc, hdrs, tbl_rows)
                        doc.add_paragraph("")
        else:
            doc.add_paragraph(
                "No improved tool definitions were generated."
            )

    # ==================================================================
    #  Verification
    # ==================================================================

    def _verification(self, doc: Document):
        doc.add_paragraph(
            "Final before/after comparison. This section is the "
            "qualification evidence for the changes applied."
        )

        headers = ["Metric", "Before", "After", "Delta", "Status"]
        rows: list[dict] = []

        if self.pr:
            d = self.pr.score_after - self.pr.score_before
            rows.append({
                "Metric": "Prompt Score",
                "Before": _pct(self.pr.score_before),
                "After": _pct(self.pr.score_after),
                "Delta": f"{'+' if d >= 0 else ''}{_pct(d)}",
                "Status": (
                    "IMPROVED" if d > 0
                    else "SAME" if d == 0 else "REGRESSED"
                ),
            })
            plen_before = len(self.pr.original_prompt)
            plen_after = (
                len(self.pr.improved_prompt)
                if self.pr.improved_prompt else 0
            )
            rows.append({
                "Metric": "Prompt Length",
                "Before": f"{plen_before} chars",
                "After": (
                    f"{plen_after} chars"
                    if self.pr.improved_prompt else "\u2014"
                ),
                "Delta": (
                    f"{plen_after - plen_before:+d}"
                    if self.pr.improved_prompt else "\u2014"
                ),
                "Status": "OK",
            })

        if self.tr:
            d = self.tr.score_after - self.tr.score_before
            rows.append({
                "Metric": "Tool Pipeline Score",
                "Before": _pct(self.tr.score_before),
                "After": _pct(self.tr.score_after),
                "Delta": f"{'+' if d >= 0 else ''}{_pct(d)}",
                "Status": (
                    "IMPROVED" if d > 0
                    else "SAME" if d == 0 else "REGRESSED"
                ),
            })
            rows.append({
                "Metric": "Schema Issues",
                "Before": str(len(self.tr.schema_issues)),
                "After": "\u2014",
                "Delta": "\u2014",
                "Status": "REVIEWED",
            })
            rows.append({
                "Metric": "Test Pass Rate",
                "Before": f"{self.tr.test_pass_rate:.0%}",
                "After": "\u2014",
                "Delta": "\u2014",
                "Status": "TESTED",
            })

        _table(doc, headers, rows, color_col="Status")

    # ==================================================================
    #  Appendix
    # ==================================================================

    def _appendix(self, doc: Document):
        doc.add_heading("Methodology", level=2)
        methods = [
            (
                "LIME Perturbation",
                "Drops/masks each sentence and measures LLM output change. "
                "High Importance = output changes substantially without "
                "that sentence.",
            ),
            (
                "Counterfactual Analysis",
                "Generates alternate prompt versions (negation, removal, "
                "substitution) and tests whether the output meaning flips.",
            ),
            (
                "Concept Extraction",
                "Identifies which abstract concepts (formality, specificity, "
                "persona, constraints) are present vs. missing.",
            ),
            (
                "Schema Audit",
                "Each tool parameter is checked for type, description, "
                "enum constraints, nesting depth, and naming clarity.",
            ),
            (
                "Prompt-Tool Alignment",
                "System prompt is scanned for tool references, usage "
                "guidance, JSON format instructions, and fallback policy.",
            ),
            (
                "Live Testing",
                "Test cases are sent to the LLM; extracted tool calls are "
                "compared against expected tool + parameters.",
            ),
            (
                "Auto-Improvement",
                "LLM generates an improved prompt and tool schemas "
                "incorporating all findings, then re-tests for "
                "verification.",
            ),
        ]
        for name, desc in methods:
            p = doc.add_paragraph()
            p.add_run(f"{name}: ").bold = True
            p.add_run(desc)

        doc.add_heading("Matrix Definitions", level=2)
        defs: list[tuple[str, str]] = []
        if self.pr:
            defs.extend([
                (
                    "M1 \u2014 PCAM",
                    "Prompt Component Assessment Matrix. Rates each "
                    "sentence\u2019s importance and fragility.",
                ),
                (
                    "M2 \u2014 SRM",
                    "Sensitivity & Robustness Matrix. FMEA-style "
                    "perturbation impact analysis.",
                ),
                (
                    "M3 \u2014 RGAM",
                    "Requirements Gap Analysis. Desired behavior vs. "
                    "prompt coverage.",
                ),
                (
                    "M4 \u2014 CAR",
                    "Corrective Action Register. Every finding with "
                    "severity and fix.",
                ),
                (
                    "M5 \u2014 BAVM",
                    "Before/After Verification Matrix. Qualification "
                    "test results.",
                ),
            ])
        if self.tr:
            defs.extend([
                (
                    "T1 \u2014 TSQM",
                    "Tool Schema Quality Matrix. ICD-style parameter "
                    "audit.",
                ),
                (
                    "T2 \u2014 PTAM",
                    "Prompt-Tool Alignment Matrix. "
                    "Requirements-to-interface traceability.",
                ),
                (
                    "T3 \u2014 TCAM",
                    "Tool Call Accuracy Matrix. Test procedure pass/fail "
                    "results.",
                ),
                (
                    "T4 \u2014 JCM",
                    "JSON Compliance Matrix. Structured output "
                    "validation.",
                ),
                (
                    "T5 \u2014 CAR",
                    "Tool Corrective Action Register. Prioritized "
                    "findings with fixes.",
                ),
            ])
        for name, desc in defs:
            p = doc.add_paragraph()
            p.add_run(f"{name}: ").bold = True
            p.add_run(desc)

        doc.add_paragraph("")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("\u2014 End of Report \u2014")
        run.font.color.rgb = _GRAY
        run.font.size = Pt(10)
        run.italic = True


# Backward-compat alias so older code still works
ToolWordReport = WordReport
